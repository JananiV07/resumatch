import pytest
from docx import Document
from pypdf import PdfWriter
from resumatch.parser import parse, ParseError


def test_parse_txt():
    text = parse("resume.txt", "Hello  world\n\n".encode("utf-8"))
    assert text == "Hello world"


def test_parse_txt_latin1_fallback():
    text = parse("resume.txt", b"caf\xe9")
    assert "caf" in text


def test_parse_docx(tmp_path):
    p = tmp_path / "r.docx"
    doc = Document()
    doc.add_paragraph("Senior Python Engineer")
    doc.add_paragraph("5 years experience")
    doc.save(p)
    text = parse("r.docx", p.read_bytes())
    assert "Senior Python Engineer" in text
    assert "5 years experience" in text


def test_parse_pdf(tmp_path):
    p = tmp_path / "r.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with open(p, "wb") as f:
        writer.write(f)
    text = parse("r.pdf", p.read_bytes())
    assert isinstance(text, str)


def test_unsupported_extension_raises():
    with pytest.raises(ParseError):
        parse("image.png", b"\x89PNG")


def test_corrupt_pdf_raises_parseerror():
    with pytest.raises(ParseError):
        parse("broken.pdf", b"not a real pdf")
